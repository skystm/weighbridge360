from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Estilos base ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
    doc.add_paragraph()
    return table

# ===================== CAPA =====================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WeighBridge 360')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Handoff — Sessão 1 → Sessão 2')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SLC Agrícola · Spark Sessions · 28 de julho de 2026')
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Este documento é o pacote de contexto oficial da Sessão 1.\nA Sessão 2 começa aqui.')
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_page_break()

# ===================== RESUMO DA SESSÃO =====================
add_heading(doc, '1. Resumo da Sessão 1', 1)
add_table(doc,
    ['Campo', 'Valor'],
    [
        ['Data', '28 de julho de 2026'],
        ['Duração', '~2h30 (13h00 → 15h30)'],
        ['Modo', 'Comprimido — problema bem definido pelo time, fases de ideação aceleradas'],
        ['Artefatos gerados', '9 de 9 ✅'],
    ]
)

# ===================== O PROBLEMA =====================
add_heading(doc, '2. O Problema (acordado pelo time)', 1)

add_para(doc, 'HMW (How Might We):', bold=True)
p = doc.add_paragraph(style='Quote')
p.add_run('"Como podemos garantir o registro seguro e em tempo real do peso de produtos e insumos agrícolas — totalmente integrado ao SAP — para prevenir fraudes, reduzir perdas financeiras e aumentar a visibilidade operacional do campo ao ERP?"')

add_para(doc, 'Dor central:', bold=True)
add_para(doc, 'O processo de pesagem nas fazendas da SLC Agrícola é inteiramente manual — porteiro digita placa, balanceiro digita pesos, fiscal recebe e-mail para emitir DANFe, coordenador monitora por WhatsApp. O sistema legado cai 465 vezes/mês sem fallback. Nenhum dado chega ao SAP automaticamente. O custo é R$ 4,1 milhões/mês em atrasos e risco real de fraude sem rastreabilidade.')

# ===================== SOLUÇÃO =====================
add_heading(doc, '3. A Solução Escolhida', 1)

add_para(doc, 'Nome da solução: WeighBridge 360', bold=True)
p = doc.add_paragraph(style='Quote')
p.add_run('"Plataforma SAP-native que automatiza o fluxo de pesagem agrícola — do campo ao ERP — eliminando trabalho manual, prevenindo fraudes e entregando visibilidade em tempo real para operadores e gestores."')

add_para(doc, 'Arquitetura de alto nível:', bold=True)
add_table(doc,
    ['Camada', 'Componente', 'Responsabilidade'],
    [
        ['Operação', 'WeighBridge 360 App (SAP BTP)', 'OCR de placa, captura automática de peso, emissão DANFe/CTe, modo offline'],
        ['Integração', 'SAP Integration Suite (CPI)', 'Transporte confiável com retry automático para o SAP HANA'],
        ['Persistência', 'SAP HANA', 'Banco transacional — fonte de verdade dos tickets'],
        ['Governança', 'SAP Datasphere', 'Consolidação multi-fazenda, qualidade de dados'],
        ['Analytics', 'SAP Analytics Cloud (SAC)', 'Dashboard em tempo real para coordenador'],
    ]
)

add_para(doc, 'Foco de prioridade: 70% balanceiro (operação) / 30% coordenador + porteiro', bold=True)

# ===================== PERSONA =====================
add_heading(doc, '4. Persona Principal', 1)
add_para(doc, 'Ricardo Mendes Souza — Balanceiro, 34 anos, Fazenda Paiaguás (MT)', bold=True)
p = doc.add_paragraph(style='Quote')
p.add_run('"Eu faço minha parte direito, mas quando o sistema cai, nada que eu fiz antes existiu. Aí a culpa é minha."')

add_para(doc, 'O que Ricardo precisa:', bold=True)
for item in [
    'Sistema que salva automaticamente a cada etapa — nunca perde dados',
    'Captura de peso sem digitação — sensor integrado',
    'DANFe/CTe em segundos — sem e-mail, sem espera',
    'Trilha de auditoria que o protege de responsabilizações injustas',
]:
    p = doc.add_paragraph(item, style='List Bullet')

# ===================== VALUE PROPOSITION =====================
add_heading(doc, '5. Value Proposition', 1)
p = doc.add_paragraph(style='Quote')
p.add_run('"WeighBridge 360 permite que o balanceiro registre pesagens sem digitar nada e que o coordenador enxergue o campo em tempo real — eliminando retrabalho, fraude e atraso documental, economizando até R$ 2,9 milhões por mês para a SLC Agrícola."')

add_para(doc, 'Métricas de sucesso:', bold=True)
add_table(doc,
    ['Métrica', 'Hoje', 'Meta'],
    [
        ['Tempo médio de pesagem', '>50 min', '<15 min'],
        ['Campos digitados por ticket', '~12', '0'],
        ['Tempo de emissão DANFe', 'até 2h', '<30 segundos'],
        ['Sincronização SAP', 'Manual após o processo', '100% automático em tempo real'],
        ['Monitoramento do coordenador', 'WhatsApp', 'Dashboard SAC em tempo real'],
    ]
)

# ===================== DECISÕES =====================
add_heading(doc, '6. Decisões Tomadas pelo Time', 1)
add_table(doc,
    ['Decisão', 'O que foi decidido'],
    [
        ['Automação do check-in', 'OCR de placa (confiança ≥ 0.85; abaixo → confirmação manual)'],
        ['Captura de peso', 'Sensor automático; entrada manual apenas como fallback com justificativa'],
        ['Emissão fiscal', 'DANFe e CTe disparados automaticamente ao registrar Pesagem 2'],
        ['Integração SAP', 'CPI → SAP HANA → Datasphere → SAC'],
        ['Resiliência', 'Modo offline obrigatório com sincronização automática ao reconectar'],
        ['Governança', 'Trilha de auditoria imutável (append-only) em todo ticket'],
        ['Divergência de peso', 'Ticket bloqueado automaticamente → aprovação exclusiva do coordenador'],
        ['Padronização', 'Mesmo fluxo para as 25 fazendas — configuração por fazenda, não customização'],
    ]
)

# ===================== FORA DO ESCOPO =====================
add_heading(doc, '7. Fora do Escopo', 1)
for item in [
    'Automação do carregamento físico (silos, armazéns)',
    'Integração com sistemas de transportadoras externas',
    'Substituição ou migração do SAP ERP',
    'Gestão de contratos de fornecedores',
    'Outros processos agrícolas (plantio, colheita, irrigação)',
    'App mobile nativo (o app do balanceiro é web responsivo no BTP)',
]:
    doc.add_paragraph(item, style='List Bullet')

# ===================== ARTEFATOS =====================
add_heading(doc, '8. Artefatos Gerados — Sessão 1', 1)
add_table(doc,
    ['Artefato', 'Arquivo', 'Status'],
    [
        ['Problem Statement', 'projeto/sessao1/problem-statement.md', '✅'],
        ['Persona (Ricardo)', 'projeto/sessao1/persona.md', '✅'],
        ['Interview Simulation', 'projeto/sessao1/interview-simulation.md', '✅'],
        ['User Journey (As-Is / To-Be)', 'projeto/sessao1/user-journey.md', '✅'],
        ['Insights Cluster + Solução', 'projeto/sessao1/insights-cluster.md', '✅'],
        ['Value Proposition', 'projeto/sessao1/value-proposition.md', '✅'],
        ['Domain Glossary (19 termos)', 'projeto/sessao1/domain-glossary.md', '✅'],
        ['Domain Model (14 BR + 5 contextos)', 'projeto/sessao1/domain-model.md', '✅'],
        ['Handoff para Sessão 2', 'projeto/sessao1/handoff-para-sessao2.md', '✅'],
    ]
)

# ===================== GLOSSÁRIO =====================
add_heading(doc, '9. Termos Críticos para a Sessão 2', 1)
add_table(doc,
    ['Termo', 'Significado rápido'],
    [
        ['TicketPesagem', 'Unidade de rastreabilidade — um ciclo completo de pesagem'],
        ['Tara', 'Peso do caminhão vazio (Pesagem 1)'],
        ['PesoBruto', 'Peso do caminhão cheio (Pesagem 2)'],
        ['PesoLiquido', 'Bruto − Tara — sempre calculado, nunca digitado'],
        ['TrilhaAuditoria', 'Log imutável append-only de cada ação no ticket'],
        ['Divergencia', 'Status de bloqueio quando peso é inválido ou inconsistente'],
        ['ModoOffline', 'Operação local sem internet com sync automático ao reconectar'],
        ['ContratoSAP', 'Vínculo que valida placa, produto e quantidade na entrada'],
    ]
)

# ===================== PONTOS DE ATENÇÃO =====================
add_heading(doc, '10. Pontos de Atenção para a Sessão 2', 1)
atencoes = [
    ('Integração com balança física', 'O protocolo de comunicação com o sensor precisa ser confirmado com a SLC (RS-232, TCP/IP, Modbus?). Marcar como dependência de hardware na especificação.'),
    ('Integração SEFAZ para DANFe/CTe', 'Emissão automática depende de certificado digital da SLC no ambiente SEFAZ. Verificar se já está configurado no SAP.'),
    ('Contratos SAP como fonte de dados', 'Definir como o sistema busca contratos ativos — API SAP, OData ou replicação para HANA.'),
    ('Modo offline — escopo do protótipo', 'Definir o que será implementado vs. simulado no pitch. Sugestão: offline real para Pesagem 1 e 2; sincronização demonstrável.'),
    ('Aprovação de divergência', 'Fluxo de divergência é diferencial de governança para o pitch — implementar mesmo que simplificado.'),
]
for titulo, descricao in atencoes:
    p = doc.add_paragraph()
    run = p.add_run(f'{titulo}: ')
    run.bold = True
    p.add_run(descricao)

# ===================== PRÓXIMOS PASSOS =====================
add_heading(doc, '11. Próximos Passos — Sessão 2', 1)
add_table(doc,
    ['Prioridade', 'Artefato', 'Agente'],
    [
        ['1', 'User Stories com critérios de aceite', 'NEXUS'],
        ['2', 'API Contracts (endpoints principais)', 'NEXUS'],
        ['3', 'Especificação técnica (product spec)', 'NEXUS + PRISM'],
        ['4', 'Test Scenarios + Validation Matrix', 'FORGE'],
        ['5', 'Revisão S2', 'LENS'],
    ]
)

add_para(doc, 'Para iniciar a Sessão 2, diga "sessão 2" ou "vamos especificar".', italic=True)

# ===================== RODAPÉ =====================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WeighBridge 360 · SLC Agrícola · Spark Sessions · Sessão 1 de 3')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

output_path = r'C:\Users\I757723\Desktop\IW26\projeto\sessao1\WeighBridge360-Handoff-Sessao1.docx'
doc.save(output_path)
print(f'Arquivo salvo: {output_path}')
